"""附件测试：Schema 校验 / 路径安全 / 能力门槛 / 双协议 content parts / 节点接线。"""
import base64
import json

import pytest

from aps.schemas.attachments import Attachment, AttachmentList
from aps.services import attachments as att_svc
from aps.services.adapters.chat_adapter import ChatCompletionsAdapter, _attachment_parts
from aps.services.adapters.responses_adapter import _attachment_input_items
from aps.services.gateway import Gateway, GenerateRequest


# ---------------------------------------------------------------- Schema

def test_attachment_from_text():
    a = Attachment.from_text("hello world", name="note.txt")
    assert a.kind == "text"
    assert a.size_bytes == len("hello world".encode())
    assert a.content == "hello world"
    assert a.validate() == []


def test_attachment_from_data_uri_image():
    b64 = base64.b64encode(b"\x89PNG-fake").decode()
    a = Attachment.from_data_uri(f"data:image/png;base64,{b64}", name="pic.png")
    assert a.kind == "image"
    assert a.mime_type == "image/png"
    assert a.is_data_uri is True
    assert a.size_bytes == 9


def test_attachment_name_strips_path():
    a = Attachment.from_text("x", name="../../etc/passwd")
    assert a.name == "passwd"          # 只保留展示名，防路径语义


def test_attachment_size_limits():
    big = Attachment(kind="text", name="big.txt", size_bytes=1024 * 1024 * 2)
    assert any("上限" in p for p in big.validate())
    ok = Attachment(kind="text", name="ok.txt", size_bytes=100)
    assert ok.validate() == []


# ---------------------------------------------------------------- 路径安全

def test_load_path_attachments_text_and_image(tmp_path):
    text_file = tmp_path / "doc.txt"
    text_file.write_text("第一行\n第二行", encoding="utf-8")
    img_file = tmp_path / "pic.png"
    img_file.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 16)
    atts, warnings = att_svc.load_path_attachments(
        f"{text_file.name}\n{img_file.name}", base_dir=str(tmp_path))
    assert warnings == []
    by_kind = {a.kind: a for a in atts}
    assert by_kind["text"].content.replace("\r\n", "\n") == "第一行\n第二行"
    assert by_kind["image"].kind == "image"
    assert by_kind["image"].as_data_uri().startswith("data:image/png;base64,")


def test_load_path_attachments_rejects_traversal(tmp_path):
    evil = tmp_path / ".." / "outside.txt"
    evil.parent.mkdir(exist_ok=True)
    evil.write_text("secret", encoding="utf-8")
    atts, warnings = att_svc.load_path_attachments(
        "../outside.txt", base_dir=str(tmp_path))
    assert atts == []
    assert any("越界" in w for w in warnings)


def test_load_path_attachments_rejects_absolute(tmp_path):
    target = tmp_path / "abs.txt"
    target.write_text("x", encoding="utf-8")
    atts, warnings = att_svc.load_path_attachments(
        str(target), base_dir=str(tmp_path / "sub"))
    assert atts == []
    assert warnings


def test_load_path_attachments_skips_missing_and_oversize(tmp_path):
    atts, warnings = att_svc.load_path_attachments("nope.txt",
                                                   base_dir=str(tmp_path))
    assert atts == []
    assert any("跳过" in w for w in warnings)


# ---------------------------------------------------------------- 能力门槛

def test_gate_text_always_ok():
    a = Attachment.from_text("hi")
    sendable, warnings, err = att_svc.gate_attachments([a], {}, False, False)
    assert sendable and err is None


def test_gate_image_without_vision_errors():
    img = Attachment.from_data_uri("data:image/png;base64," + base64.b64encode(b"x").decode())
    _, _, err = att_svc.gate_attachments([img], {"vision": False}, False, False)
    assert err and "视觉" in err and "supports_vision" in err


def test_gate_image_with_supports_vision_ok():
    img = Attachment.from_data_uri("data:image/png;base64," + base64.b64encode(b"x").decode())
    sendable, _, err = att_svc.gate_attachments([img], {"vision": False}, True, False)
    assert sendable and err is None


def test_gate_file_without_files_errors():
    f = Attachment(kind="file", name="a.bin", mime_type="application/octet-stream",
                   content="data:application/octet-stream;base64,AA==",
                   is_data_uri=True, size_bytes=1)
    _, _, err = att_svc.gate_attachments([f], {"files": False}, False, False)
    assert err and "supports_files" in err


def test_gate_mixed_reports_all_errors():
    img = Attachment.from_data_uri("data:image/png;base64," + base64.b64encode(b"x").decode())
    _, _, err = att_svc.gate_attachments([img, img], {}, False, False)
    assert err and "无法发送" in err


def _minimal_pdf(text: str) -> bytes:
    """构造一个含单行文本的最小合法 PDF（带正确 xref 偏移，pypdf 可解析）。"""
    stream = f"BT /F1 12 Tf 20 100 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] "
         b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"),
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = b"%PDF-1.4\n"
    offsets = []
    for i, obj in enumerate(objects, 1):
        offsets.append(len(out))
        out += b"%d 0 obj\n%s\nendobj\n" % (i, obj)
    xref_pos = len(out)
    out += b"xref\n0 %d\n" % (len(objects) + 1)
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += (b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n"
            % (len(objects) + 1, xref_pos))
    return out


# ---------------------------------------------------------------- 本地文档提取（P0-5）

def test_local_extract_pdf_text():
    pdf = _minimal_pdf("Hello from pdf file")
    att = Attachment.from_base64(base64.b64encode(pdf).decode(),
                                 name="doc.pdf", mime_type="application/pdf")
    assert att.kind == "file"
    extracted = att_svc.local_extract_document(att)
    assert extracted is not None
    text_att, note = extracted
    assert text_att.kind == "text"
    assert "Hello from pdf file" in text_att.content
    assert note == ""


def test_local_extract_docx_text(tmp_path):
    import docx as _docx
    doc = _docx.Document()
    doc.add_paragraph("第一段：合同条款")
    doc.add_paragraph("第二段：付款方式")
    buf = tmp_path / "d.docx"
    doc.save(str(buf))
    att = Attachment.from_base64(base64.b64encode(buf.read_bytes()).decode(),
                                 name="d.docx",
                                 mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    extracted = att_svc.local_extract_document(att)
    assert extracted is not None
    assert "第一段" in extracted[0].content and "付款方式" in extracted[0].content


def test_local_extract_scanned_pdf_none():
    pdf = _minimal_pdf("")  # 无文本层：空 stream
    att = Attachment.from_base64(base64.b64encode(pdf).decode(),
                                 name="scan.pdf", mime_type="application/pdf")
    assert att_svc.local_extract_document(att) is None


def test_local_extract_non_document_none():
    f = Attachment(kind="file", name="a.bin", mime_type="application/octet-stream",
                   content="data:application/octet-stream;base64,AA==",
                   is_data_uri=True, size_bytes=1)
    assert att_svc.local_extract_document(f) is None


def test_gate_pdf_without_files_local_extract():
    pdf = _minimal_pdf("Terms and conditions")
    att = Attachment.from_base64(base64.b64encode(pdf).decode(),
                                 name="terms.pdf", mime_type="application/pdf")
    sendable, warnings, err = att_svc.gate_attachments(
        [att], {"files": False}, False, False)
    assert err is None
    assert sendable and sendable[0].kind == "text"
    assert any("本地提取文本发送" in w for w in warnings)


def test_gate_scanned_pdf_without_files_errors():
    pdf = _minimal_pdf("")
    att = Attachment.from_base64(base64.b64encode(pdf).decode(),
                                 name="scan.pdf", mime_type="application/pdf")
    _, _, err = att_svc.gate_attachments([att], {"files": False}, False, False)
    assert err and "无法发送" in err and "扫描" in err


def test_gateway_pdf_degrades_to_text_with_warning(monkeypatch, store):
    store.create_profile({"profile_id": "p1"})
    store.set_capabilities("p1", {"responses": True, "files": False})
    pdf = _minimal_pdf("Hello from pdf file")
    att = Attachment.from_base64(base64.b64encode(pdf).decode(),
                                 name="doc.pdf", mime_type="application/pdf")
    captured = {}

    class FakeAdapter:
        def generate(self, profile, api_key, **kw):
            captured.update(kw)
            from aps.schemas.results import LLMResult
            return LLMResult(profile_id="p1", text="ok", protocol="responses")

    gw = Gateway(store=store)
    gw._responses = FakeAdapter()
    gw._chat = FakeAdapter()
    result = gw.generate(store.get_profile("p1"), "k",
                         GenerateRequest(messages=[], attachments=[att]))
    assert not result.has_error()
    assert any("本地提取文本发送" in w for w in result.warnings)
    assert captured["attachments"] and captured["attachments"][0].kind == "text"
    assert "Hello from pdf file" in captured["attachments"][0].content


def test_gateway_binary_file_without_files_errors(store):
    store.create_profile({"profile_id": "p1"})
    store.set_capabilities("p1", {"responses": True, "files": False})
    f = Attachment(kind="file", name="a.bin", mime_type="application/octet-stream",
                   content="data:application/octet-stream;base64,AA==",
                   is_data_uri=True, size_bytes=1)
    gw = Gateway(store=store)
    result = gw.generate(store.get_profile("p1"), "k",
                         GenerateRequest(messages=[], attachments=[f]))
    assert result.has_error()
    assert result.error.kind == "attachment_unsupported"


# ---------------------------------------------------------------- 附件 warning 到节点输出（P0-4）

def test_attachment_path_warning_reaches_node_output(monkeypatch, store, tmp_path):
    import aps.nodes.llm_chat as mod
    doc = tmp_path / "ok.txt"
    doc.write_text("正常附件", encoding="utf-8")
    store.create_profile({"profile_id": "p1"})
    store.set_api_key("p1", "sk-abcdef1234567890")
    payload = store.get_profile("p1").node_payload()

    class FakeGateway:
        def generate(self, profile, api_key, req):
            from aps.schemas.results import LLMResult
            return LLMResult(text="done", profile_id="p1")

    monkeypatch.setattr(mod, "Gateway", lambda: FakeGateway())
    monkeypatch.setattr(att_svc, "default_input_dir", lambda: str(tmp_path))
    node = mod.APS_LLMGenerate()
    # 一行正常文件 + 一行越界路径：越界必须出现在 warnings 输出
    text, _, _, _, _, _, warnings = node.generate(
        AI_PROFILE=payload, system_prompt="", user_prompt="总结", context="",
        session=None, history_mode="off", output_mode="text", json_schema="",
        attachment_files=f"{doc.name}\n../outside.txt")
    assert any("已跳过" in w for w in warnings.splitlines())


def test_attachment_skipped_warning(monkeypatch, store, tmp_path):
    import aps.nodes.llm_chat as mod
    store.create_profile({"profile_id": "p1"})
    store.set_api_key("p1", "sk-abcdef1234567890")
    payload = store.get_profile("p1").node_payload()

    class FakeGateway:
        def generate(self, profile, api_key, req):
            from aps.schemas.results import LLMResult
            return LLMResult(text="done", profile_id="p1")

    monkeypatch.setattr(mod, "Gateway", lambda: FakeGateway())
    monkeypatch.setattr(att_svc, "default_input_dir", lambda: str(tmp_path))
    node = mod.APS_LLMGenerate()
    text, _, _, _, _, _, warnings = node.generate(
        AI_PROFILE=payload, system_prompt="", user_prompt="总结", context="",
        session=None, history_mode="off", output_mode="text", json_schema="",
        attachment_files="missing_file.txt")
    assert any("已跳过" in w for w in warnings.splitlines())


# ---------------------------------------------------------------- 双协议 content parts

def _b64(mime, payload):
    return f"data:{mime};base64," + base64.b64encode(payload).decode()


def test_responses_attachment_input_items():
    atts = [Attachment.from_text("见附件", name="a.txt"),
            Attachment.from_data_uri(_b64("image/png", b"png"), name="pic.png"),
            Attachment.from_base64(base64.b64encode(b"bin"), name="f.bin",
                                   mime_type="application/pdf")]
    items = _attachment_input_items(atts)
    kinds = [i["content"][0]["type"] for i in items]
    assert kinds == ["input_text", "input_image", "input_file"]
    assert items[1]["content"][0]["image_url"].startswith("data:image/png;base64,")
    assert items[2]["content"][0]["filename"] == "f.bin"


def test_chat_attachment_parts():
    atts = [Attachment.from_text("见附件"),
            Attachment.from_data_uri(_b64("image/png", b"png"), name="pic.png"),
            Attachment.from_base64(base64.b64encode(b"bin"), name="f.bin",
                                   mime_type="application/pdf")]
    parts = _attachment_parts(atts)
    assert parts[0] == {"type": "text", "text": "见附件"}
    assert parts[1]["type"] == "image_url"
    assert parts[1]["image_url"]["url"].startswith("data:image/png;base64,")
    assert parts[2]["type"] == "file"
    assert parts[2]["file"]["filename"] == "f.bin"


def test_chat_attachments_merge_into_last_user_message():
    from aps.schemas.results import ChatMessage
    msgs = [ChatMessage(role="user", content="请总结")]
    atts = [Attachment.from_text("数据")]
    parts = _attachment_parts(atts)
    content = msgs[0].content
    merged = (content if isinstance(content, list)
              else [{"type": "text", "text": content}]) + parts
    assert merged[1] == {"type": "text", "text": "数据"}


# ---------------------------------------------------------------- 网关/节点

def test_gateway_attachment_unsupported_returns_error(store):
    store.create_profile({"profile_id": "p1"})
    store.set_capabilities("p1", {"vision": False, "files": False})
    img = Attachment.from_data_uri(_b64("image/png", b"png"), name="pic.png")
    gw = Gateway(store=store)
    result = gw.generate(store.get_profile("p1"), "k", GenerateRequest(
        messages=[], attachments=[img]))
    assert result.has_error()
    assert result.error.kind == "attachment_unsupported"


def test_gateway_attachment_supported_reaches_adapter(monkeypatch, store):
    store.create_profile({"profile_id": "p1", "supports_vision": True,
                          "supports_files": True})
    store.set_capabilities("p1", {"responses": True, "vision": False})
    img = Attachment.from_data_uri(_b64("image/png", b"png"), name="pic.png")
    calls = {}

    class FakeAdapter:
        def generate(self, profile, api_key, **kw):
            calls.update(kw)
            from aps.schemas.results import LLMResult
            return LLMResult(profile_id="p1", text="ok", protocol="responses")

    gw = Gateway(store=store)
    gw._responses = FakeAdapter()
    gw._chat = FakeAdapter()
    result = gw.generate(store.get_profile("p1"), "k",
                         GenerateRequest(messages=[], attachments=[img]))
    assert not result.has_error()
    assert calls["attachments"] and calls["attachments"][0].kind == "image"


def test_llm_chat_node_attachment_files(monkeypatch, store, tmp_path):
    import aps.nodes.llm_chat as mod
    doc = tmp_path / "doc.txt"
    doc.write_text("文件内容", encoding="utf-8")
    store.create_profile({"profile_id": "p1", "name": "D",
                          "supports_vision": True, "supports_files": True})
    store.set_api_key("p1", "sk-abcdef1234567890")
    payload = store.get_profile("p1").node_payload()

    captured = {}

    class FakeGateway:
        def generate(self, profile, api_key, req):
            captured["req"] = req
            from aps.schemas.results import LLMResult
            return LLMResult(text="done", profile_id="p1")

    monkeypatch.setattr(mod, "Gateway", lambda: FakeGateway())
    # 附件路径解析基目录指向临时目录
    monkeypatch.setattr(att_svc, "default_input_dir", lambda: str(tmp_path))
    node = mod.APS_LLMGenerate()
    node.generate(AI_PROFILE=payload, system_prompt="", user_prompt="总结",
                  context="", session=None, history_mode="off",
                  output_mode="text", json_schema="",
                  attachment_files=f"{doc.name}")
    assert captured["req"].attachments
    assert captured["req"].attachments[0].kind == "text"
    assert captured["req"].attachments[0].content == "文件内容"
