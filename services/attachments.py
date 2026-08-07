"""附件服务：本地路径安全解析 + 能力门槛（视觉/文件）与降级规则。

安全（§32）：路径必须解析在受控输入目录内（拒绝 .. 逃逸 / 绝对路径绕过）；
内容不写日志；大小上限；附件生命周期短（文本/图片直接内联，无临时文件遗留）。
降级规则（产品决策 D20；0.2.1 补充）：
- 文本附件：任何协议直接作为文本内容发送，无需能力；
- 图片附件：主模型必须支持视觉（caps.vision 或档案 supports_vision），否则**报错**，
  绝不静默丢弃伪装成功；
- 文件附件：支持文件发送（caps.files 或档案 supports_files）→ 发送文件内容部分；
  否则 PDF/DOCX → 本地轻量提取文本（pypdf/python-docx，可选依赖）→ 文本降级 + warning；
  提取失败（扫描件无文本层/依赖缺失/其他二进制）→ 报错，绝不假装识别。
"""
from __future__ import annotations

import base64
import mimetypes
import os
from pathlib import Path
from typing import List, Optional, Tuple

from ..schemas.attachments import (
    MAX_FILE_BYTES,
    MAX_IMAGE_BYTES,
    MAX_TEXT_BYTES,
    Attachment,
)

_TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv", ".yml", ".yaml", ".log",
                    ".srt", ".vtt", ".xml", ".py", ".js", ".ts", ".html", ".css"}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".avif"}
# 可本地提取文本的文档类型（0.2.1：Provider 无文件能力时的降级路径）
_DOC_EXTENSIONS = {".pdf", ".docx"}


def default_input_dir() -> Optional[str]:
    """ComfyUI input 目录（真实环境经 folder_paths；测试环境回退到 cwd）。"""
    import os
    env = os.environ.get("COMFYUI_INPUT_DIR")
    if env and os.path.isdir(env):
        return env
    try:
        import folder_paths
        return folder_paths.get_input_directory()
    except Exception:  # noqa: BLE001 - 非 ComfyUI 环境
        return os.getcwd()


def _resolve_safe(raw_path: str, base_dir: Optional[str]) -> Optional[Path]:
    """把用户给的路径安全解析到 base_dir（ComfyUI input 目录）内。

    规则：相对路径以 base_dir 为基准解析；绝对路径/.. 逃逸后必须仍位于
    base_dir 之下；拒绝任何解析到 base 之外的路径。
    """
    raw = (raw_path or "").strip().strip('"')
    if not raw:
        return None
    base = Path(base_dir).resolve() if base_dir else None
    raw_obj = Path(raw).expanduser()
    if os.path.isabs(raw) or raw_obj.is_absolute():
        candidate = raw_obj.resolve()
    elif base is not None:
        candidate = (base / raw_obj).resolve()
    else:
        candidate = raw_obj.resolve()
    if base is not None:
        try:
            candidate.relative_to(base)
        except ValueError:
            return None
    if not candidate.is_file():
        return None
    return candidate


def _read_attachment(path: Path) -> Attachment:
    size = path.stat().st_size
    ext = path.suffix.lower()
    name = path.name
    mime, _ = mimetypes.guess_type(name)
    if ext in _TEXT_EXTENSIONS or (mime or "").startswith("text/"):
        if size > MAX_TEXT_BYTES:
            return Attachment(kind="text", name=name, mime_type=mime or "text/plain",
                              size_bytes=size,
                              problems=[f"文本附件超过大小上限（{MAX_TEXT_BYTES} 字节）"])
        data = path.read_bytes()
        return Attachment(kind="text", name=name, mime_type=mime or "text/plain",
                          content=data.decode("utf-8", errors="replace"),
                          size_bytes=size, source=f"file:{name}")
    if ext in _IMAGE_EXTENSIONS or (mime or "").startswith("image/"):
        if size > MAX_IMAGE_BYTES:
            return Attachment(kind="image", name=name, mime_type=mime or "image/png",
                              size_bytes=size,
                              problems=[f"图片附件超过大小上限（{MAX_IMAGE_BYTES} 字节）"])
        b64 = base64.b64encode(path.read_bytes()).decode("ascii")
        return Attachment.from_base64(b64, name=name,
                                      mime_type=mime or "image/png")
    # 其他文件：作为文件附件（发送 file 内容部分；无能力时按降级规则处理）
    if size > MAX_FILE_BYTES:
        return Attachment(kind="file", name=name, mime_type=mime or "application/octet-stream",
                          size_bytes=size,
                          problems=[f"文件附件超过大小上限（{MAX_FILE_BYTES} 字节）"])
    b64 = base64.b64encode(path.read_bytes()).decode("ascii")
    return Attachment.from_base64(b64, name=name,
                                  mime_type=mime or "application/octet-stream")


def _document_extractable(name: str, mime_type: str) -> bool:
    """判断该附件是否为可本地提取文本的文档（PDF/DOCX）。"""
    mime = (mime_type or "").lower()
    if mime in ("application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
        return True
    ext = (name or "").rsplit(".", 1)[-1].lower() if "." in (name or "") else ""
    return ext in (".pdf", ".docx")


def _extract_text_from_payload(att: Attachment) -> str:
    """从附件 base64 载荷中本地提取文本（PDF/DOCX）。

    不 OCR：扫描 PDF 没有文本层时返回空串（由调用方明确报错）。
    依赖 pypdf / python-docx 为可选；未安装时抛 ImportError（调用方转为可读错误）。
    """
    payload = att.content or ""
    if att.is_data_uri and "," in payload:
        payload = payload.split(",", 1)[1]
    try:
        raw = base64.b64decode(payload)
    except Exception:  # noqa: BLE001 - 非法 base64
        return ""
    from io import BytesIO

    name = att.name or ""
    if name.lower().endswith(".pdf"):
        from pypdf import PdfReader  # noqa: PLC0415 - 可选依赖，延迟导入

        reader = PdfReader(BytesIO(raw))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if name.lower().endswith(".docx"):
        import docx  # noqa: PLC0415 - 可选依赖，延迟导入

        document = docx.Document(BytesIO(raw))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    return ""


def local_extract_document(att: Attachment) -> Optional[Tuple[Attachment, str]]:
    """Provider 无文件能力时：把 PDF/DOCX 本地提取为文本附件。

    返回 (文本附件, 附加说明) 或 None（无法提取：非文档/扫描件无文本层/依赖缺失）。
    提取文本超 MAX_TEXT_BYTES 时截断并在说明中明确标注（不静默丢弃）。
    """
    if not _document_extractable(att.name, att.mime_type):
        return None
    try:
        text = _extract_text_from_payload(att)
    except ImportError as exc:
        raise ValueError(
            f"本地提取 {att.name} 需要 pypdf/python-docx（可选依赖，未安装）：{exc}") from exc
    except Exception:  # noqa: BLE001 - 解析失败按无法提取处理
        return None
    if not text or not text.strip():
        return None  # 扫描件无文本层，不假装识别
    note = ""
    if len(text.encode("utf-8")) > MAX_TEXT_BYTES:
        text = text[:MAX_TEXT_BYTES]
        note = f"（提取文本已截断至 {MAX_TEXT_BYTES} 字节）"
    out = Attachment(kind="text", name=att.name, mime_type="text/plain",
                     content=text, size_bytes=len(text.encode("utf-8")),
                     source=f"local_extract:{att.name}")
    return out, note


def load_path_attachments(paths_text: str,
                          base_dir: Optional[str] = None) -> Tuple[List[Attachment], List[str]]:
    """把多行文件路径解析为附件列表（安全）。返回 (attachments, warnings)。"""
    attachments: List[Attachment] = []
    warnings: List[str] = []
    for line in (paths_text or "").splitlines():
        path = _resolve_safe(line, base_dir)
        if path is None:
            warnings.append(f"附件路径不可访问或越界，已跳过：{line.strip()!r}")
            continue
        att = _read_attachment(path)
        if att.problems:
            warnings.append(f"附件 {att.name} 校验失败：{'；'.join(att.problems)}，已跳过")
            continue
        attachments.append(att)
    return attachments, warnings


def gate_attachments(attachments: List[Attachment], caps: dict,
                     supports_vision: bool, supports_files: bool) -> Tuple[List[Attachment], List[str], Optional[str]]:
    """能力门槛：返回 (可发送列表, warnings, error_text)。

    有图片/文件附件无法发送时返回 error_text（调用方应报错，绝不静默成功）。
    """
    if not attachments:
        return [], [], None
    sendable: List[Attachment] = []
    warnings: List[str] = []
    errors: List[str] = []
    vision_ok = bool(supports_vision or caps.get("vision") is True)
    files_ok = bool(supports_files or caps.get("files") is True)

    for a in attachments:
        if a.kind == "text":
            sendable.append(a)
        elif a.kind == "image":
            if vision_ok:
                sendable.append(a)
            else:
                errors.append(
                    f"图片附件 {a.name!r} 无法发送：当前档案不支持视觉"
                    "（可配置视觉档案或开启档案高级设置 supports_vision）")
        elif a.kind == "file":
            if files_ok:
                sendable.append(a)
            else:
                # 降级：PDF/DOCX 本地提取文本（0.2.1）；其余二进制无法提取 → 报错
                try:
                    extracted = local_extract_document(a)
                except ValueError as exc:
                    errors.append(str(exc))
                    continue
                if extracted is not None:
                    sendable.append(extracted[0])
                    warnings.append(
                        f"Provider 不支持原生文件输入，已本地提取文本发送："
                        f"{a.name}{extracted[1]}")
                else:
                    errors.append(
                        f"文件附件 {a.name!r} 无法发送：当前服务不支持文件内容部分，"
                        "且该文件无法本地提取文本（扫描件无文本层或非 PDF/DOCX；"
                        "可在档案高级设置开启 supports_files）")
    if errors:
        return sendable, warnings, "；".join(errors)
    return sendable, warnings, None


def text_context_for(attachments: List[Attachment]) -> str:
    """把文本附件拼成注入上下文的块（不进日志）。"""
    parts = []
    for a in attachments:
        if a.kind == "text" and a.content:
            parts.append(f"--- 附件 {a.name} ---\n{a.content}")
    return "\n\n".join(parts)
